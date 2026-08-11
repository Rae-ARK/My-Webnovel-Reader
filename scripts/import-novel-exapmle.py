from pathlib import Path
from html import escape
import json
import re
import sqlite3
import sys

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DB = ROOT / "public/content/library.sqlite"

DEFAULT_CONFIG_PATH = ROOT / "scripts/novel.config.json"
EXAMPLE_CONFIG_PATH = ROOT / "scripts/novel.config.example.json"

REQUIRED_CONFIG_KEYS = (
    "fiction_id",
    "title",
    "author",
    "status",
    "cover",
    "docx",
)


def load_config(path):
    if not path.exists():
        raise RuntimeError(
            f"Novel config not found at {path}.\n"
            f"Copy {EXAMPLE_CONFIG_PATH.relative_to(ROOT)} to "
            f"{path.relative_to(ROOT)} and fill in this fiction's "
            "details. This file is author-specific and is not committed."
        )

    with path.open(encoding="utf-8") as handle:
        config = json.load(handle)

    missing = [key for key in REQUIRED_CONFIG_KEYS if key not in config]

    if missing:
        raise RuntimeError(
            f"Novel config at {path} is missing required key(s): "
            f"{', '.join(missing)}"
        )

    return config


config_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_CONFIG_PATH
config = load_config(config_path)

FICTION_ID = config["fiction_id"]
TITLE = config["title"]
AUTHOR = config["author"]
STATUS = config["status"]
COVER = config["cover"]
DOCX = ROOT / config["docx"]

INDEX_PATTERN = re.compile(r"^\s*index\b", re.IGNORECASE)
CHAPTER_PATTERN = re.compile(r"^\s*chapter\s+(\d+)\b", re.IGNORECASE)


def paragraph_to_html(paragraph):
    parts = []

    for run in paragraph.runs:
        text = escape(run.text)

        if not text:
            continue

        if run.bold:
            text = f"<strong>{text}</strong>"

        if run.italic:
            text = f"<em>{text}</em>"

        parts.append(text)

    text = "".join(parts).strip()

    return f"<p>{text}</p>" if text else ""


def classify_entry(title):
    normalized = title.strip().lower()

    if normalized.startswith("interlude"):
        return "interlude"

    if normalized.startswith("extra"):
        return "extra"

    if normalized.startswith("author's afterward"):
        return "afterword"

    if normalized.startswith("afterword"):
        return "afterword"

    if normalized.startswith("afterward"):
        return "afterword"

    return "chapter"


def get_chapter_number(title):
    match = CHAPTER_PATTERN.match(title)

    if not match:
        return None

    return int(match.group(1))


doc = Document(DOCX)

heading_indexes = [
    i
    for i, paragraph in enumerate(doc.paragraphs)
    if paragraph.text.strip()
    and paragraph.style.name == "Heading 1"
]

if not heading_indexes:
    raise RuntimeError("No Heading 1 entries found")


sections = []

for position, start in enumerate(heading_indexes):
    end = (
        heading_indexes[position + 1]
        if position + 1 < len(heading_indexes)
        else len(doc.paragraphs)
    )

    title = doc.paragraphs[start].text.strip()
    content = []

    for paragraph in doc.paragraphs[start + 1:end]:
        if paragraph.style.name == "Heading 2":
            continue

        html = paragraph_to_html(paragraph)

        if html:
            content.append(html)

    sections.append({
        "title": title,
        "content": "\n".join(content),
    })


DEFAULT_INDEX_TITLE = "Index"

# Indexes and their entries are derived entirely from the document's own
# structure. A section whose title matches INDEX_PATTERN (e.g. "Index of
# ARC 1", "Part One", "Volume I", "Book 2", ...) declares a new index using
# whatever title the author wrote. Every readable entry that follows belongs
# to the most recently declared index. If the document never declares an
# index heading before its first entry, a single default index titled
# "Index" is created to hold it. Nothing here is specific to any one
# fiction's naming scheme, chapter count, or numbering.

indexes = []
entries = []

current_index = None
entry_number = 0


def start_new_index(title):
    index = {
        "title": title,
        "position": len(indexes),
        "entries": [],
    }
    indexes.append(index)
    return index


for section in sections:
    title = section["title"]

    if INDEX_PATTERN.match(title):
        current_index = start_new_index(title)
        continue

    entry_number += 1

    entry_type = classify_entry(title)
    chapter_number = (
        get_chapter_number(title) if entry_type == "chapter" else None
    )

    if entry_type == "chapter" and chapter_number is None:
        raise RuntimeError(
            f"Could not determine chapter number from: {title!r}"
        )

    entry = {
        "id": f"{FICTION_ID}-entry-{entry_number}",
        "fiction_id": FICTION_ID,
        "type": entry_type,
        "number": chapter_number,
        "title": title,
        "content": section["content"],
    }

    entries.append(entry)

    if current_index is None:
        # No explicit index heading has appeared yet in the document.
        # Every fiction still gets an index; entries encountered before
        # any explicit heading fall under this default one.
        current_index = start_new_index(DEFAULT_INDEX_TITLE)

    current_index["entries"].append(entry)


if not entries:
    raise RuntimeError("No readable entries found in document")



with sqlite3.connect(DB) as conn:
    conn.execute("PRAGMA foreign_keys = ON")

    conn.executescript("""
        DROP VIEW IF EXISTS fiction_summary;

        DROP TABLE IF EXISTS index_entries;
        DROP TABLE IF EXISTS indexes;
        DROP TABLE IF EXISTS content_entries;
        DROP TABLE IF EXISTS chapters;

        CREATE TABLE content_entries (
          id TEXT PRIMARY KEY,
          fiction_id TEXT NOT NULL,
          type TEXT NOT NULL
            CHECK (type IN (
              'chapter',
              'interlude',
              'extra',
              'afterword'
            )),
          number INTEGER,
          title TEXT NOT NULL,
          content TEXT NOT NULL,
          FOREIGN KEY (fiction_id)
            REFERENCES fictions(id)
            ON DELETE CASCADE
        );

        CREATE INDEX idx_content_entries_fiction
          ON content_entries (fiction_id);

        CREATE INDEX idx_content_entries_fiction_number
          ON content_entries (fiction_id, number);

        CREATE TABLE indexes (
          id TEXT PRIMARY KEY,
          fiction_id TEXT NOT NULL,
          title TEXT NOT NULL,
          position INTEGER NOT NULL,
          FOREIGN KEY (fiction_id)
            REFERENCES fictions(id)
            ON DELETE CASCADE,
          UNIQUE (fiction_id, position)
        );

        CREATE INDEX idx_indexes_fiction_position
          ON indexes (fiction_id, position);

        CREATE TABLE index_entries (
          index_id TEXT NOT NULL,
          entry_id TEXT NOT NULL,
          position INTEGER NOT NULL,
          label TEXT,
          PRIMARY KEY (index_id, entry_id),
          FOREIGN KEY (index_id)
            REFERENCES indexes(id)
            ON DELETE CASCADE,
          FOREIGN KEY (entry_id)
            REFERENCES content_entries(id)
            ON DELETE CASCADE,
          UNIQUE (index_id, position)
        );

        CREATE INDEX idx_index_entries_index_position
          ON index_entries (index_id, position);

        CREATE VIEW fiction_summary AS
        SELECT
          f.id,
          f.title,
          f.author,
          f.cover,
          f.synopsis,
          f.status,
          COUNT(c.id) AS entry_count
        FROM fictions f
        LEFT JOIN content_entries c
          ON c.fiction_id = f.id
        GROUP BY f.id;
    """)

    conn.execute(
        "DELETE FROM fiction_genres WHERE fiction_id = ?",
        (FICTION_ID,),
    )

    conn.execute(
        "DELETE FROM fiction_tags WHERE fiction_id = ?",
        (FICTION_ID,),
    )

    conn.execute(
        "DELETE FROM fictions WHERE id = ?",
        (FICTION_ID,),
    )

    conn.execute(
        """
        INSERT INTO fictions
        (id, title, author, cover, synopsis, status)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (
            FICTION_ID,
            TITLE,
            AUTHOR,
            COVER,
            TITLE,
            STATUS,
        ),
    )

    conn.executemany(
        """
        INSERT INTO content_entries
        (id, fiction_id, type, number, title, content)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        [
            (
                entry["id"],
                entry["fiction_id"],
                entry["type"],
                entry["number"],
                entry["title"],
                entry["content"],
            )
            for entry in entries
        ],
    )

    for index in indexes:
        index_id = (
            f"{FICTION_ID}-index-{index['position'] + 1}"
        )

        conn.execute(
            """
            INSERT INTO indexes
            (id, fiction_id, title, position)
            VALUES (?, ?, ?, ?)
            """,
            (
                index_id,
                FICTION_ID,
                index["title"],
                index["position"],
            ),
        )

        conn.executemany(
            """
            INSERT INTO index_entries
            (index_id, entry_id, position, label)
            VALUES (?, ?, ?, ?)
            """,
            [
                (
                    index_id,
                    entry["id"],
                    position,
                    entry["title"],
                )
                for position, entry
                in enumerate(index["entries"])
            ],
        )

    conn.commit()


print("IMPORT OK")
print(f"Title: {TITLE}")
print(f"Author: {AUTHOR}")
print(f"Readable entries: {len(entries)}")
print(f"Indexes: {len(indexes)}")

for index in indexes:
    print(
        f"  {index['position'] + 1}. "
        f"{index['title']}: "
        f"{len(index['entries'])} entries"
    )
