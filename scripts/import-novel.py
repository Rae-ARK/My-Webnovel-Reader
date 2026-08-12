"""
Import a fiction from a .docx source document into library.sqlite,
per a novel.config.json file (see novel.config.example.json).

Safe to re-run against a database that already contains other
fictions -- only the fiction named by the config's fiction_id is ever
touched. Safe to re-run against a fiction that was already imported,
as long as you tell it what to do about the conflict:

    --overwrite     Fully replace the existing fiction with this import.
    --merge         Keep the existing fiction; add new entries, and
                     resolve entries that already exist per --on-conflict.
    --on-conflict {skip,overwrite,abort}
                     How to resolve entries that already exist during a
                     --merge. Default: prompt interactively, or abort if
                     --non-interactive.

Neither flag given, and the fiction already exists: the script refuses
and explains the two options (a "merge conflict" error), rather than
silently overwriting or silently doing nothing.

Flags for scripted/CI use:
    --non-interactive   Never prompt; every decision must come from a flag.
    --yes / -y           Skip the "are you sure?" confirmation for --overwrite.
    --dry-run             Parse and report what would happen; write nothing.
    --db PATH             Target database path (default: public/content/library.sqlite).

Examples:
    # First import of a new fiction.
    python3 scripts/import-novel.py

    # Re-import after editing the source doc, keeping existing entries
    # that weren't touched and overwriting ones that were:
    python3 scripts/import-novel.py --merge --on-conflict overwrite

    # Scripted full replace, e.g. in a deploy pipeline:
    python3 scripts/import-novel.py --overwrite --yes --non-interactive --db public/content/library.sqlite
"""

import argparse
from pathlib import Path
from html import escape
import json
import re
import shutil
import sqlite3
import sys

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DB = ROOT / "public/content/library.sqlite"

DEFAULT_CONFIG_PATH = ROOT / "scripts/novel.config.json"
EXAMPLE_CONFIG_PATH = ROOT / "scripts/novel.config.example.json"

REQUIRED_CONFIG_KEYS = (
    "fiction_id",
    "title",
    "author",
    "status",
    "cover",
    "docx",
    "synopsis",
)


def parse_args():
    parser = argparse.ArgumentParser(
        description="Import a fiction from a .docx source into library.sqlite.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument(
        "config",
        nargs="?",
        type=Path,
        default=DEFAULT_CONFIG_PATH,
        help="Path to novel.config.json (default: scripts/novel.config.json)",
    )
    parser.add_argument("--db", type=Path, default=DEFAULT_DB, help="Target database path")
    conflict_group = parser.add_mutually_exclusive_group()
    conflict_group.add_argument(
        "--overwrite",
        action="store_true",
        help="Fully replace an existing fiction with this import",
    )
    conflict_group.add_argument(
        "--merge",
        action="store_true",
        help="Keep an existing fiction; add/update entries per --on-conflict",
    )
    parser.add_argument(
        "--on-conflict",
        choices=("skip", "overwrite", "abort"),
        default=None,
        help="How to resolve entries that already exist during --merge",
    )
    parser.add_argument(
        "--non-interactive",
        action="store_true",
        help="Never prompt; every decision must come from a flag",
    )
    parser.add_argument("-y", "--yes", action="store_true", help="Skip confirmation prompts")
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Parse and report what would happen; write nothing",
    )
    return parser.parse_args()


def load_config(path):
    if not path.exists():
        raise RuntimeError(
            f"Novel config not found at {path}.\n"
            f"Copy {EXAMPLE_CONFIG_PATH.relative_to(ROOT)} to "
            f"{path.relative_to(ROOT)} and fill in this fiction's "
            "details. This file is author-specific and is not committed."
        )

    with path.open(encoding="utf-8") as handle:
        config = json.load(handle)

    missing = [key for key in REQUIRED_CONFIG_KEYS if key not in config]

    if missing:
        raise RuntimeError(
            f"Novel config at {path} is missing required key(s): "
            f"{', '.join(missing)}"
        )

    return config


args = parse_args()
config_path = args.config
DB = args.db
config = load_config(config_path)

FICTION_ID = config["fiction_id"]
TITLE = config["title"]
AUTHOR = config["author"]
STATUS = config["status"]
COVER = config["cover"]
DOCX = ROOT / config["docx"]
SYNOPSIS = config["synopsis"]

INDEX_PATTERN = re.compile(r"^\s*index\b", re.IGNORECASE)
CHAPTER_PATTERN = re.compile(r"^\s*chapter\s+(\d+)\b", re.IGNORECASE)


def paragraph_to_html(paragraph):
    parts = []

    for run in paragraph.runs:
        text = escape(run.text)

        if not text:
            continue

        if run.bold:
            text = f"<strong>{text}</strong>"

        if run.italic:
            text = f"<em>{text}</em>"

        parts.append(text)

    text = "".join(parts).strip()

    return f"<p>{text}</p>" if text else ""


def classify_entry(title):
    normalized = title.strip().lower()

    if normalized.startswith("interlude"):
        return "interlude"

    if normalized.startswith("extra"):
        return "extra"

    if normalized.startswith("author's afterward"):
        return "afterword"

    if normalized.startswith("afterword"):
        return "afterword"

    if normalized.startswith("afterward"):
        return "afterword"

    return "chapter"


def get_chapter_number(title):
    match = CHAPTER_PATTERN.match(title)

    if not match:
        return None

    return int(match.group(1))


doc = Document(DOCX)

heading_indexes = [
    i
    for i, paragraph in enumerate(doc.paragraphs)
    if paragraph.text.strip()
    and paragraph.style.name == "Heading 1"
]

if not heading_indexes:
    raise RuntimeError("No Heading 1 entries found")


sections = []

for position, start in enumerate(heading_indexes):
    end = (
        heading_indexes[position + 1]
        if position + 1 < len(heading_indexes)
        else len(doc.paragraphs)
    )

    title = doc.paragraphs[start].text.strip()
    content = []
    raw_lines = []

    for paragraph in doc.paragraphs[start + 1:end]:
        if paragraph.style.name == "Heading 2":
            continue

        text = paragraph.text.strip()

        if text:
            raw_lines.append(text)

        html = paragraph_to_html(paragraph)

        if html:
            content.append(html)

    sections.append({
        "title": title,
        "content": "\n".join(content),
        "raw_lines": raw_lines,
        "is_index_heading": bool(INDEX_PATTERN.match(title)),
    })


DEFAULT_INDEX_TITLE = "Index"

# Indexes are derived entirely from the document's own structure. A
# section whose title matches INDEX_PATTERN (e.g. "Index of ARC 1",
# "Part One", "Volume I", "Book 2", ...) declares a new index using
# whatever title the author wrote. Nothing here is specific to any one
# fiction's naming scheme, chapter count, or numbering.
#
# Two document layouts are supported, both driven purely by content the
# author already wrote -- never by fiction-specific code:
#
#   1. Interleaved: each index heading is immediately followed by the
#      entries that belong to it, then the next index heading. Entries
#      are simply grouped under whichever index heading most recently
#      appeared in document order.
#
#   2. Table-of-contents block: every index heading (with a manifest of
#      chapter numbers under it) appears up front, and the real,
#      readable entries all follow afterward. Grouping entries by
#      "most recently declared heading" would put everything under the
#      last index in this layout, so instead each index's own manifest
#      is scanned for chapter numbers, giving it a numeric range (e.g.
#      an index whose manifest mentions chapters 6-13 owns numbers
#      6 through 13). Real entries are then bucketed by matching their
#      own chapter number against these ranges. Entries with no number
#      (interludes, extras, afterwords) join whichever index most
#      recently claimed a numbered chapter, so they land next to the
#      chapter they follow in the reading order.
#
# A fiction whose index headings carry no discoverable numbers at all
# (no manifest, or non-numeric organization) automatically falls back
# to the interleaved behaviour, so both layouts are handled by the same
# generic pass.

indexes = []
entries = []


def start_new_index(title, max_chapter_number):
    index = {
        "title": title,
        "position": len(indexes),
        "entries": [],
        "max_chapter_number": max_chapter_number,
        "range_start": None,
        "range_end": None,
    }
    indexes.append(index)
    return index


# First pass: declare every index in document order, and, if its own
# listing mentions chapter numbers, record the highest one as that
# index's manifest ceiling.
for section in sections:
    if not section["is_index_heading"]:
        continue

    mentioned_numbers = [
        int(match.group(1))
        for line in section["raw_lines"]
        for match in [CHAPTER_PATTERN.match(line)]
        if match
    ]

    max_chapter_number = max(mentioned_numbers) if mentioned_numbers else None

    start_new_index(section["title"], max_chapter_number)

# Turn each index's ceiling into a contiguous numeric range. An index
# with no ceiling (no numbers found in its own listing) gets no range
# and is skipped entirely by number-based lookup below.
previous_ceiling = 0

for index in indexes:
    if index["max_chapter_number"] is None:
        continue

    index["range_start"] = previous_ceiling + 1
    index["range_end"] = index["max_chapter_number"]
    previous_ceiling = index["max_chapter_number"]


def index_for_chapter_number(number):
    for index in indexes:
        if index["range_start"] is None:
            continue

        if index["range_start"] <= number <= index["range_end"]:
            return index

    return None


# Second pass: walk the document again in order, building readable
# entries and bucketing each one into an index.
index_iterator = iter(indexes)
current_declared_index = None
last_numbered_bucket = None
entry_number = 0

for section in sections:
    title = section["title"]

    if section["is_index_heading"]:
        current_declared_index = next(index_iterator)
        continue

    entry_number += 1

    entry_type = classify_entry(title)
    chapter_number = (
        get_chapter_number(title) if entry_type == "chapter" else None
    )

    if entry_type == "chapter" and chapter_number is None:
        raise RuntimeError(
            f"Could not determine chapter number from: {title!r}"
        )

    entry = {
        "id": f"{FICTION_ID}-entry-{entry_number}",
        "fiction_id": FICTION_ID,
        "type": entry_type,
        "number": chapter_number,
        "title": title,
        "content": section["content"],
    }

    entries.append(entry)

    bucket = None

    if chapter_number is not None:
        bucket = index_for_chapter_number(chapter_number)

        if bucket is not None:
            last_numbered_bucket = bucket

    if bucket is None:
        # No numeric range claimed this entry (either it has no number,
        # or no index declared a manifest range at all) -- fall back to
        # whichever index most recently owned a numbered chapter, then
        # to whichever index heading was most recently declared in
        # document order (the interleaved-layout behaviour).
        bucket = last_numbered_bucket or current_declared_index

    if bucket is None:
        # No index heading has appeared yet in the document at all.
        # Every fiction still gets an index; entries encountered before
        # any explicit heading fall under this default one.
        bucket = start_new_index(DEFAULT_INDEX_TITLE, None)
        current_declared_index = bucket

    bucket["entries"].append(entry)


if not entries:
    raise RuntimeError("No readable entries found in document")



SCHEMA_DDL = """
    CREATE TABLE IF NOT EXISTS fictions (
      id TEXT PRIMARY KEY,
      title TEXT NOT NULL,
      author TEXT NOT NULL,
      cover TEXT,
      synopsis TEXT NOT NULL,
      status TEXT NOT NULL CHECK (status IN ('ongoing', 'completed', 'hiatus'))
    );

    CREATE TABLE IF NOT EXISTS genres (
      id TEXT PRIMARY KEY,
      name TEXT NOT NULL UNIQUE
    );

    CREATE TABLE IF NOT EXISTS fiction_genres (
      fiction_id TEXT NOT NULL,
      genre_id TEXT NOT NULL,
      PRIMARY KEY (fiction_id, genre_id),
      FOREIGN KEY (fiction_id) REFERENCES fictions(id) ON DELETE CASCADE,
      FOREIGN KEY (genre_id) REFERENCES genres(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS tags (
      id TEXT PRIMARY KEY,
      name TEXT NOT NULL UNIQUE
    );

    CREATE TABLE IF NOT EXISTS fiction_tags (
      fiction_id TEXT NOT NULL,
      tag_id TEXT NOT NULL,
      PRIMARY KEY (fiction_id, tag_id),
      FOREIGN KEY (fiction_id) REFERENCES fictions(id) ON DELETE CASCADE,
      FOREIGN KEY (tag_id) REFERENCES tags(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS content_entries (
      id TEXT PRIMARY KEY,
      fiction_id TEXT NOT NULL,
      type TEXT NOT NULL
        CHECK (type IN (
          'chapter',
          'interlude',
          'extra',
          'afterword'
        )),
      number INTEGER,
      title TEXT NOT NULL,
      content TEXT NOT NULL,
      FOREIGN KEY (fiction_id)
        REFERENCES fictions(id)
        ON DELETE CASCADE
    );

    CREATE INDEX IF NOT EXISTS idx_content_entries_fiction
      ON content_entries (fiction_id);

    CREATE INDEX IF NOT EXISTS idx_content_entries_fiction_number
      ON content_entries (fiction_id, number);

    CREATE TABLE IF NOT EXISTS indexes (
      id TEXT PRIMARY KEY,
      fiction_id TEXT NOT NULL,
      title TEXT NOT NULL,
      position INTEGER NOT NULL,
      FOREIGN KEY (fiction_id)
        REFERENCES fictions(id)
        ON DELETE CASCADE,
      UNIQUE (fiction_id, position)
    );

    CREATE INDEX IF NOT EXISTS idx_indexes_fiction_position
      ON indexes (fiction_id, position);

    CREATE TABLE IF NOT EXISTS index_entries (
      index_id TEXT NOT NULL,
      entry_id TEXT NOT NULL,
      position INTEGER NOT NULL,
      label TEXT,
      PRIMARY KEY (index_id, entry_id),
      FOREIGN KEY (index_id)
        REFERENCES indexes(id)
        ON DELETE CASCADE,
      FOREIGN KEY (entry_id)
        REFERENCES content_entries(id)
        ON DELETE CASCADE,
      UNIQUE (index_id, position)
    );

    CREATE INDEX IF NOT EXISTS idx_index_entries_index_position
      ON index_entries (index_id, position);
"""

# The view has to be dropped/recreated on every run since its
# definition can change across versions of this script, but that's
# harmless -- a view holds no data of its own.
VIEW_DDL = """
    DROP VIEW IF EXISTS fiction_summary;

    CREATE VIEW fiction_summary AS
    SELECT
      f.id,
      f.title,
      f.author,
      f.cover,
      f.synopsis,
      f.status,
      COUNT(c.id) AS entry_count
    FROM fictions f
    LEFT JOIN content_entries c
      ON c.fiction_id = f.id
    GROUP BY f.id;
"""


def wipe_fiction_content(conn, fiction_id):
    """Delete everything belonging to one fiction. Scoped to
    fiction_id throughout -- never touches other fictions' rows."""
    conn.execute(
        """
        DELETE FROM index_entries
        WHERE index_id IN (SELECT id FROM indexes WHERE fiction_id = ?)
        """,
        (fiction_id,),
    )
    conn.execute("DELETE FROM indexes WHERE fiction_id = ?", (fiction_id,))
    conn.execute("DELETE FROM content_entries WHERE fiction_id = ?", (fiction_id,))
    conn.execute("DELETE FROM fiction_genres WHERE fiction_id = ?", (fiction_id,))
    conn.execute("DELETE FROM fiction_tags WHERE fiction_id = ?", (fiction_id,))
    conn.execute("DELETE FROM fictions WHERE id = ?", (fiction_id,))


def insert_fiction_content(conn, fiction_id, entries_to_insert):
    conn.execute(
        """
        INSERT INTO fictions
        (id, title, author, cover, synopsis, status)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (fiction_id, TITLE, AUTHOR, COVER, SYNOPSIS, STATUS),
    )

    conn.executemany(
        """
        INSERT INTO content_entries
        (id, fiction_id, type, number, title, content)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        [
            (e["id"], e["fiction_id"], e["type"], e["number"], e["title"], e["content"])
            for e in entries_to_insert
        ],
    )

    write_indexes(conn, fiction_id)


def write_indexes(conn, fiction_id):
    """(Re)writes every index/index_entries row for this fiction from
    the freshly parsed document structure. Always a full rebuild of
    the table of contents, independent of whether individual entries
    were skipped/overwritten/inserted -- their ids already exist in
    content_entries by the time this runs."""
    for index in indexes:
        index_id = f"{fiction_id}-index-{index['position'] + 1}"

        conn.execute(
            "INSERT INTO indexes (id, fiction_id, title, position) VALUES (?, ?, ?, ?)",
            (index_id, fiction_id, index["title"], index["position"]),
        )

        conn.executemany(
            """
            INSERT INTO index_entries
            (index_id, entry_id, position, label)
            VALUES (?, ?, ?, ?)
            """,
            [
                (index_id, entry["id"], position, entry["title"])
                for position, entry in enumerate(index["entries"])
            ],
        )


def resolve_conflicts(conflicting_titles, args):
    if args.on_conflict:
        return args.on_conflict

    if args.non_interactive:
        raise RuntimeError(
            f"{len(conflicting_titles)} entr"
            f"{'y' if len(conflicting_titles) == 1 else 'ies'} already "
            "exist and --non-interactive was given. Pass "
            "--on-conflict {skip,overwrite,abort} to resolve this without prompting."
        )

    print(f"{len(conflicting_titles)} entries already exist in this fiction:")
    for title in conflicting_titles[:10]:
        print(f"  - {title}")
    if len(conflicting_titles) > 10:
        print(f"  ... and {len(conflicting_titles) - 10} more")

    while True:
        choice = input("Resolve as [s]kip / [o]verwrite / [a]bort? ").strip().lower()
        if choice in ("s", "skip"):
            return "skip"
        if choice in ("o", "overwrite"):
            return "overwrite"
        if choice in ("a", "abort"):
            return "abort"
        print("Please answer skip, overwrite, or abort.")


with sqlite3.connect(DB) as conn:
    conn.execute("PRAGMA foreign_keys = ON")
    conn.executescript(SCHEMA_DDL)

    existing = conn.execute(
        "SELECT title, author FROM fictions WHERE id = ?", (FICTION_ID,)
    ).fetchone()

    if existing is None:
        # Brand new fiction -- --overwrite/--merge are irrelevant here.
        plan = "insert"
    elif args.overwrite:
        plan = "overwrite"
    elif args.merge:
        plan = "merge"
    else:
        existing_title, existing_author = existing
        raise RuntimeError(
            f"Fiction '{FICTION_ID}' already exists in the database "
            f"(title: {existing_title!r}, author: {existing_author!r}) -- "
            "merge conflict.\n"
            "Pass --overwrite to fully replace it with this import, or "
            "--merge to add/update its content without deleting entries "
            "this import doesn't touch."
        )

    if plan == "overwrite" and existing is not None and not args.yes:
        if args.non_interactive:
            raise RuntimeError(
                "--overwrite on an existing fiction requires --yes when "
                "--non-interactive is set (refusing to overwrite silently)."
            )
        existing_title, _ = existing
        confirm = input(
            f"Overwrite existing fiction '{FICTION_ID}' ({existing_title!r})? [y/N] "
        ).strip().lower()
        if confirm not in ("y", "yes"):
            raise RuntimeError("Aborted by user.")

    conflict_resolution = None
    conflicting_ids = set()

    if plan == "merge":
        existing_ids = {
            row[0]
            for row in conn.execute(
                "SELECT id FROM content_entries WHERE fiction_id = ?", (FICTION_ID,)
            )
        }
        new_ids = {e["id"] for e in entries}
        conflicting_ids = existing_ids & new_ids

        if conflicting_ids:
            conflicting_titles = [e["title"] for e in entries if e["id"] in conflicting_ids]
            conflict_resolution = resolve_conflicts(conflicting_titles, args)

            if conflict_resolution == "abort":
                raise RuntimeError(
                    f"Aborted: {len(conflicting_ids)} entries already exist. "
                    "Re-run with --on-conflict skip or --on-conflict overwrite."
                )

    if args.dry_run:
        print("DRY RUN — no changes written")
        print(f"Plan: {plan}")
        print(f"Title: {TITLE}")
        print(f"Author: {AUTHOR}")
        print(f"Readable entries parsed: {len(entries)}")
        print(f"Indexes parsed: {len(indexes)}")
        if plan == "merge":
            new_count = len(entries) - len(conflicting_ids)
            print(f"  would insert: {new_count} new entries")
            if conflicting_ids:
                print(
                    f"  would {conflict_resolution}: {len(conflicting_ids)} "
                    "already-existing entries"
                )
        sys.exit(0)

    if DB.exists():
        backup_path = DB.with_name(DB.name + ".bak")
        shutil.copy2(DB, backup_path)
        print(f"Existing database found — backed up to {backup_path}")

    conn.executescript(VIEW_DDL)

    if plan in ("insert", "overwrite"):
        if existing is not None:
            wipe_fiction_content(conn, FICTION_ID)
        insert_fiction_content(conn, FICTION_ID, entries)
    else:  # merge
        conn.execute(
            """
            UPDATE fictions
            SET title = ?, author = ?, cover = ?, synopsis = ?, status = ?
            WHERE id = ?
            """,
            (TITLE, AUTHOR, COVER, SYNOPSIS, STATUS, FICTION_ID),
        )

        to_insert = [e for e in entries if e["id"] not in conflicting_ids]
        to_overwrite = (
            [e for e in entries if e["id"] in conflicting_ids]
            if conflict_resolution == "overwrite"
            else []
        )

        conn.executemany(
            """
            INSERT INTO content_entries
            (id, fiction_id, type, number, title, content)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            [
                (e["id"], e["fiction_id"], e["type"], e["number"], e["title"], e["content"])
                for e in to_insert
            ],
        )

        conn.executemany(
            """
            UPDATE content_entries
            SET type = ?, number = ?, title = ?, content = ?
            WHERE id = ?
            """,
            [
                (e["type"], e["number"], e["title"], e["content"], e["id"])
                for e in to_overwrite
            ],
        )

        # Table of contents is always rebuilt fresh from the parsed
        # document -- every entry it references already exists in
        # content_entries by this point (inserted, overwritten, or
        # left alone as a skip).
        conn.execute(
            """
            DELETE FROM index_entries
            WHERE index_id IN (SELECT id FROM indexes WHERE fiction_id = ?)
            """,
            (FICTION_ID,),
        )
        conn.execute("DELETE FROM indexes WHERE fiction_id = ?", (FICTION_ID,))
        write_indexes(conn, FICTION_ID)

    conn.commit()


print("IMPORT OK")
print(f"Plan: {plan}")
print(f"Title: {TITLE}")
print(f"Author: {AUTHOR}")
print(f"Readable entries: {len(entries)}")
if plan == "merge":
    print(f"  new: {len(entries) - len(conflicting_ids)}")
    if conflicting_ids:
        print(f"  {conflict_resolution}: {len(conflicting_ids)}")
print(f"Indexes: {len(indexes)}")

for index in indexes:
    print(
        f"  {index['position'] + 1}. "
        f"{index['title']}: "
        f"{len(index['entries'])} entries"
    )
