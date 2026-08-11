"""
Read-only diagnostic. Prints paragraph style + a short text preview for
every paragraph in the source docx, from the start through the first
*Heading 1* paragraph whose text starts with "chapter" (i.e. the real
chapter-1 section heading, not a table-of-contents mention of it).
Does not touch the database.

Run from anywhere inside the project (project root or scripts/):
    python3 diagnose_structure.py
"""
from pathlib import Path
from docx import Document
import json
import sys

SCRIPT_DIR = Path(__file__).resolve().parent

candidates = [
    SCRIPT_DIR / "scripts" / "novel.config.json",
    SCRIPT_DIR / "novel.config.json",
    SCRIPT_DIR.parent / "scripts" / "novel.config.json",
]

config_path = next((p for p in candidates if p.exists()), None)

if config_path is None:
    print("Could not find novel.config.json. Looked in:")
    for p in candidates:
        print(f"  {p}")
    sys.exit(1)

ROOT = config_path.parents[1]

config = json.loads(config_path.read_text(encoding="utf-8"))
doc = Document(ROOT / config["docx"])

seen_real_chapter_heading = False

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name

    if text:
        print(f"{i:4} [{style:12}] {text[:70]!r}")

    if style == "Heading 1" and text.lower().startswith("chapter"):
        seen_real_chapter_heading = True
        # print a few more paragraphs of the real chapter's body so we
        # can see the content style too, then stop
        for j in range(i + 1, min(i + 4, len(doc.paragraphs))):
            t2 = doc.paragraphs[j].text.strip()
            s2 = doc.paragraphs[j].style.name
            if t2:
                print(f"{j:4} [{s2:12}] {t2[:70]!r}")
        break
